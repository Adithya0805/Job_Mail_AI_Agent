# Resume parser service using pdfplumber, python-docx, and Llama 3/Gemini model
import io
import re
import json
import pdfplumber
import docx
from services.gemini_service import GeminiService

class ScanDetectedError(Exception):
    pass

class ParsingError(Exception):
    pass

class ResumeParserService:
    def __init__(self):
        self.llm_service = GeminiService()

    def extract_text(self, file_bytes: bytes, file_type: str) -> str:
        text = ""
        file_type = file_type.lower().strip()
        
        if file_type == "pdf" or file_type == "application/pdf":
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
        elif file_type in ["docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text]
            table_cells = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text not in table_cells:
                            table_cells.append(cell.text)
            text = "\n".join(paragraphs + table_cells)
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

        # Clean excessive spaces but keep line breaks
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        return text.strip()

    def build_extraction_prompt(self, resume_text: str) -> str:
        return f"""You are an expert resume parser. Extract structured information from this resume text. Be precise — only extract information that is explicitly stated, never infer or fabricate.

Resume text:
{resume_text}

Extract and return ONLY valid JSON matching this exact structure:
{{
  "full_name": "string",
  "email": "string",
  "phone": "string",
  "location": "string",
  "linkedin": "string or null",
  "github": "string or null",
  "portfolio": "string or null",
  "degree": "string",
  "institution": "string",
  "graduation_year": number or null,
  "cgpa": number or null,
  "experience_level": "fresher | junior | mid",
  "skills": {{
    "languages": ["string"],
    "frameworks": ["string"],
    "ai_ml": ["string"]
  }},
  "projects": [
    {{
      "name": "string",
      "description": "string",
      "tech_stack": ["string"],
      "live_url": "string or null"
    }}
  ],
  "certifications": ["string"],
  "summary": "string"
}}

Rules:
- If a field is not found in the resume, use null (not empty string)
- For skills, categorize into languages/frameworks/ai_ml based on common knowledge (Python = language, React = framework, LangChain = ai_ml)
- For experience_level: 0-1 years = fresher, 1-3 years = junior, 3+ years = mid
- For projects, extract up to 5 most relevant/recent ones
- Put any resume content that doesn't fit the schema (awards, languages spoken, hobbies, etc.) into unmapped_text
- Return ONLY the JSON object, no explanation, no markdown code blocks"""

    def calculate_confidence(self, profile: dict) -> str:
        # Check high conditions
        has_name = bool(profile.get("full_name"))
        has_email = bool(profile.get("email"))
        
        skills = profile.get("skills") or {}
        languages = skills.get("languages") or []
        frameworks = skills.get("frameworks") or []
        ai_ml = skills.get("ai_ml") or []
        total_skills = len(languages) + len(frameworks) + len(ai_ml)
        
        projects = profile.get("projects") or []
        
        if has_name and has_email and total_skills >= 3 and len(projects) >= 1:
            return "high"
        elif has_name or has_email or total_skills >= 1:
            return "medium"
        else:
            return "low"

    def parse_resume(self, file_bytes: bytes, file_type: str) -> dict:
        # Step 1: Text extraction
        text = self.extract_text(file_bytes, file_type)
        
        # Step 2: Scan validation
        if len(text) < 100:
            raise ScanDetectedError("This looks like a scanned document. Please upload a text-based PDF or fill in manually.")
            
        # Step 3: Call LLM with retry loop
        prompt = self.build_extraction_prompt(text)
        retries = 2
        last_error = None
        
        for attempt in range(retries):
            try:
                # Use GeminiService's dual-provider fallbacks
                parsed_json = self.llm_service._call_llm_json(prompt, temperature=0.1)
                
                # Check for schema compatibility and cleanup fields
                if not isinstance(parsed_json, dict):
                    raise ValueError("LLM response did not parse as a JSON object.")
                
                # Calculate confidence score
                confidence = self.calculate_confidence(parsed_json)
                
                # Extract extra content from unmapped_text if present in LLM response
                unmapped_text = parsed_json.get("unmapped_text") or ""
                
                # Return result matching requirements
                return {
                    "status": "parsed",
                    "confidence": confidence,
                    "profile": parsed_json,
                    "unmapped_text": unmapped_text
                }
            except Exception as e:
                last_error = e
                # Stricter retry prompt on failure
                prompt = f"{prompt}\n\nERROR ON LAST TRY: {str(e)}\nMake sure to return valid JSON matching the exact schema specified. No preamble."
        
        raise ParsingError(f"Failed to parse resume after {retries} attempts. Error: {str(last_error)}")
